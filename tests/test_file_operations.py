from editor_window import MegasolidEditor  # Import klasy MegasolidEditor
from docx import Document  # Import do obsługi plików DOCX
from file_operations import open_txt_file, save_txt_file, open_docx_file, save_docx_file
from unittest.mock import patch
from dialogs import dialog_critical

def test_open_txt_file(qtbot, tmp_path):
    """
    Testuje otwieranie pliku TXT.
    """
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("Hello, Mordeczko!")  # Tworzymy plik testowy

    editor = MegasolidEditor()  # Teraz klasa jest zaimportowana
    qtbot.addWidget(editor)

    with patch.object(dialog_critical, "__call__") as mocked_dialog:
        open_txt_file(editor, str(txt_file))

        # Sprawdzamy, czy tekst został wczytany
        assert editor.editor.toPlainText() == "Hello, Mordeczko!"
        mocked_dialog.assert_not_called()

def test_save_txt_file(qtbot, tmp_path):
    """
    Testuje zapisywanie pliku TXT.
    """
    editor = MegasolidEditor()
    qtbot.addWidget(editor)
    editor.editor.setPlainText("Save this, Mordeczko!")

    txt_file = tmp_path / "test_save.txt"
    save_txt_file(editor, str(txt_file))

    # Sprawdzamy, czy plik został zapisany
    assert txt_file.exists()
    assert txt_file.read_text() == "Save this, Mordeczko!"

def test_open_docx_file(qtbot, tmp_path):
    """
    Testuje otwieranie pliku DOCX.
    """
    docx_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello, Mordeczko!")
    doc.save(str(docx_file))  # Tworzymy plik testowy DOCX

    editor = MegasolidEditor()
    qtbot.addWidget(editor)

    with patch.object(dialog_critical, "__call__") as mocked_dialog:
        open_docx_file(editor, str(docx_file))

        # Sprawdzamy, czy tekst został wczytany
        assert editor.editor.toPlainText() == "Hello, Mordeczko!"
        mocked_dialog.assert_not_called()

def test_save_docx_file(qtbot, tmp_path):
    """
    Testuje zapisywanie pliku DOCX.
    """
    editor = MegasolidEditor()
    qtbot.addWidget(editor)
    editor.editor.setPlainText("Save this, Mordeczko!")

    docx_file = tmp_path / "test_save.docx"
    save_docx_file(editor, str(docx_file))

    # Sprawdzamy, czy plik został zapisany
    assert docx_file.exists()

    # Otwieramy plik, żeby sprawdzić treść
    doc = Document(str(docx_file))
    text = "\n".join([para.text for para in doc.paragraphs])
    assert text == "Save this, Mordeczko!"
