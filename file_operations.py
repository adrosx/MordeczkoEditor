from constants import SUPPORTED_EXTENSIONS
from dialogs import dialog_critical, open_file_dialog, save_file_dialog
from helpers import splitext, convert_doc_to_txt, notify_user_about_conversion
from docx import Document
from PyQt6.QtWidgets import QMessageBox

def open_file(editor):
    """
    Obsługuje otwieranie plików z obsługą konwersji DOC do TXT.
    """
    path = open_file_dialog(editor)
    if not path:
        return

    ext = splitext(path)[1].lower()
    if ext == ".docx" or ext == ".doc":
        notify_user_about_conversion()
        converted_path = convert_doc_to_txt(path)
        if converted_path:
            open_txt_file(editor, converted_path)
    elif ext == ".txt":
        open_txt_file(editor, path)
    else:
        dialog_critical(editor, "Unsupported file format!")

def save_file(editor):
    """
    Obsługuje zapis pliku w bieżącej lokalizacji.
    """
    if not editor.path:
        return save_file_as(editor)

    ext = splitext(editor.path)[1].lower()
    if ext == ".docx":
        save_docx_file(editor, editor.path)
    elif ext == ".txt":
        save_txt_file(editor, editor.path)
    else:
        dialog_critical(editor, "Unsupported file format!")

def save_file_as(editor):
    """
    Obsługuje zapis pliku pod nową nazwą lub lokalizacją.
    """
    path = save_file_dialog(editor)
    if not path:
        return

    ext = splitext(path)[1].lower()
    if ext == ".docx":
        save_docx_file(editor, path)
    elif ext == ".txt":
        save_txt_file(editor, path)
    else:
        dialog_critical(editor, "Unsupported file format!")

def open_txt_file(editor, path):
    """
    Otwiera plik tekstowy i ładuje jego zawartość do edytora.
    """
    try:
        with open(path, "r", encoding="utf-8") as f:
            editor.editor.setPlainText(f.read())
        editor.path = path
        editor.update_title()
    except Exception as e:
        dialog_critical(editor, f"Error opening TXT file: {e}")

def save_txt_file(editor, path):
    """
    Zapisuje zawartość edytora do pliku tekstowego.
    """
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(editor.editor.toPlainText())
        editor.path = path
        editor.update_title()
    except Exception as e:
        dialog_critical(editor, f"Error saving TXT file: {e}")

def open_docx_file(editor, path):
    """
    Otwiera plik DOCX i ładuje jego zawartość do edytora.
    """
    try:
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        editor.editor.setPlainText(text)
        editor.path = path
        editor.update_title()
    except Exception as e:
        dialog_critical(editor, f"Error opening DOCX file: {e}")

def save_docx_file(editor, path):
    """
    Zapisuje zawartość edytora do pliku DOCX.
    """
    try:
        doc = Document()
        text = editor.editor.toPlainText()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(path)
        editor.path = path
        editor.update_title()
    except Exception as e:
        dialog_critical(editor, f"Error saving DOCX file: {e}")
